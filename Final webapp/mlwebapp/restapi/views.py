# from django.shortcuts import render
from django.http import JsonResponse,Http404
from rest_framework import status
import cv2
from fer import FER
from PIL import Image
from .models import Restmlimage,Restmlwebcam
from django.core.files.storage import FileSystemStorage
from pathlib import Path
BASE_DIR = Path(__file__).resolve().parent.parent
from PIL import Image
import numpy as np
from mtcnn import MTCNN
import pickle
import json
import tensorflow
from tensorflow.keras.layers import Dense,Input, Flatten, MaxPool2D
from tensorflow.keras.layers import Conv2D,Dropout
from tensorflow.keras import Sequential
from tensorflow.keras.models import Sequential, Model
from tensorflow.keras.regularizers import l2,l1
from tensorflow.keras.optimizers import Adam
from imutils.video import VideoStream
import imutils
from django.http.response import StreamingHttpResponse
from ipware import get_client_ip
import random
import datetime
from rest_framework.views import APIView
from rest_framework.response import Response
from .serializers import mlimageserializer

import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
from heapq import nlargest
from pdf2docxnogui import Converter
from docx import Document

# Create your views here.
class Image_detect(APIView):

    def get(self,request):
        m = Restmlimage.objects.values('id','image_upload','smile_percentage').order_by('-created_at')
        serializer = mlimageserializer(m,many=True)
        print("enters-----",m)
        return Response({'gallery':serializer.data})

    def post(self,request):
        # print("------------------",json.loads(request.body.decode('utf-8')),"----------------")
        if 'data' in request.FILES:
            print("data---------------",request.FILES['data'])
            imgs = request.FILES['data']
            a=str(imgs).split('.')
            imext = ['jpg','jpeg','png','JPEG','JPG','PNG','JFIF','jfif']
            if str(a[1]) in imext:
                im = Image.open(imgs)
                im.save('/home/ubuntu/img/'+str(imgs))
                input_image = cv2.imread('/home/ubuntu/img/'+str(imgs))
                frame_flip = input_image
                emotion_detector = FER(mtcnn=True)
                result = emotion_detector.detect_emotions(frame_flip)
                if result != []:
                    list_smile = []
                    for i in result:
                        bounding_box = i["box"]
                        emotions = i["emotions"]
                        cv2.rectangle(input_image, (
                            bounding_box[0], bounding_box[1]), (
                                          bounding_box[0] + bounding_box[2], bounding_box[1] + bounding_box[3]),
                                      (0, 155, 255), 2, )
                        # max_emotions = max(i['emotions'], key=lambda x: i['emotions'][x])
                        all_values = emotions.values()
                        max_value = max(all_values)
                        # print(max_emotions)
                        # print(max_value)
                        emotion_name, score = emotion_detector.top_emotion(frame_flip)
                        a = list(emotions.items())
                        # emotion_name = a[3][0]
                        # print(emotion_name)
                        score = a[3][1]
                        # print("score",score)
                        emotion_name = "Smile"
                        color = (255, 50, 50)
                        content = score * 100
                        if content >= 0 and content <= 20:
                            data = "Smile please,...Smile while you stil have teeth!."
                        elif content >= 21 and content <= 40:
                            data = "Smile ,it increases your face value."
                        elif content >= 41 and content <= 60:
                            data = "Your smile is literally the cutest thing,I have ever seen"
                        elif content >= 61 and content <= 80:
                            data = "Wear a smile on everyday,have a great confidence"
                        else:
                            data = "Your such a happiest and adorable one,don't forget to smile at any situation"
                        emotion_score = "{}: {}".format(emotion_name, "{:.0%}".format(score))
                        smile_percent = "{:.0%}".format(score)
                        list_smile.append(smile_percent)
                        # print("smile",smile_percent)
                        # print(smile_percent,"this is the smile percentage")
                        cv2.putText(frame_flip, emotion_score,
                                    (bounding_box[0], bounding_box[1] + bounding_box[3] + 30 + 3 * 0),
                                    cv2.FONT_HERSHEY_SIMPLEX, 0.75, color, 2, cv2.LINE_AA, )
                        # Save the result in new image file
                        cv2.imwrite(str(BASE_DIR) + "/static/detectimg/" + str(imgs), frame_flip)
                    # print(list_smile,"smile percentaages ")
                    strings = " "
                    commas = ","
                    for i in list_smile:
                        # print(i)
                        strings += str(i)
                        strings += commas
                    request.data['image_upload'] = str(imgs)
                    request.data['smile_percentage'] = strings
                    # m = Restmlimage.objects.create(image_upload = str(imgs),smile_percentage=int(50))
                    print('datas--------', request.data)
                    serializer = mlimageserializer(data=request.data)
                    if serializer.is_valid():
                        serializer.save()
                        context={'msg':serializer.data,'other_msg':data}
                        return Response(context, status=status.HTTP_201_CREATED)
                    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
                else:
                    return Response({"message":"Image doesn't have face or smile on face"},status=status.HTTP_404_NOT_FOUND)
            else:
                return Response({'message':"Upload only images"},status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response({'msg':'please upload image for smile detections'},status=status.HTTP_400_BAD_REQUEST)


class Text_file(APIView):
    def __init__(self):
        self.i = random.randint(1111,2222)

    def get(self,request):
        return Response({'msg':'Files are not saved, you can only upload and view the Abstract of your pdf. Once you upload other file or once you leave the abstract wont be saved'},status=status.HTTP_204_NO_CONTENT)

    def process(self,text):
        stopwords = list(STOP_WORDS)
        nlp = spacy.load('en_core_web_sm')
        doc = nlp(text)
        tokens = [token.text for token in doc]
        punctution = punctuation + '/n'
        word_frequencies = {}
        for word in doc:
            if word.text.lower() not in stopwords:
                if word.text.lower() not in punctution:
                    if word.text not in word_frequencies.keys():
                        word_frequencies[word.text] = 1
                    else:
                        word_frequencies[word.text] += 1
        max_frequency = max(word_frequencies.values())
        for word in word_frequencies.keys():
            word_frequencies[word] = word_frequencies[word] / max_frequency
        sentence_tokens = [sent for sent in doc.sents]
        sentence_scores = {}
        for sent in sentence_tokens:
            for word in sent:
                if word.text.lower() in word_frequencies.keys():
                    if sent not in sentence_scores.keys():
                        sentence_scores[sent] = word_frequencies[word.text.lower()]
                    else:
                        sentence_scores[sent] += word_frequencies[word.text.lower()]
        select_length = int(len(sentence_tokens) * 0.1)
        summary = nlargest(select_length, sentence_scores, key=sentence_scores.get)
        final_summary = [word.text for word in summary]
        summary = ' '.join(final_summary)
        f = open(str(BASE_DIR)+'/static/'+str(self.i)+'.txt','w')
        f.write(summary)
        f.close()
        return summary

    def process2(self,text):
        from gensim.summarization.summarizer import summarize
        abstract = summarize(text)
        return abstract

    def best(self,text):
        main_file = text
        spacytest = self.process(text)
        gensimtest = self.process2(text)

        doc1_tokens = set(main_file.lower().split())
        doc2_tokens = set(str(spacytest).lower().split())
        doc3_tokens = set(str(gensimtest).lower().split())
        
        j1 = len(doc1_tokens.intersection(doc2_tokens)) / len(doc1_tokens.union(doc2_tokens))
        j2 = len(doc1_tokens.intersection(doc3_tokens)) / len(doc1_tokens.union(doc3_tokens))
        j1 = j1 * 100
        j2 = j2 * 100
        
        context = {'file': str(self.i) + '.txt','abstract_1': str(spacytest),'a1_percentage':str(j1) ,'abstract_2':str(gensimtest), 'a2_percentage':str(j2)}
        return context

    def post(self,request):
        if 'data' in request.FILES:
            from django.core.files.storage import default_storage
            from django.core.files.base import ContentFile
            files = request.FILES['data']
            filename_split = str(files).split('.')
            file_ext = ['docx', 'pdf', 'txt']
            characters = 'abcdefghijklmNOPQRSTUVWXYZ'
            randomchar = [random.choice(characters) for i in range(10)]
            rand = ''.join(randomchar)
            if str(filename_split[1]) in file_ext:
                if str(filename_split[1]) == 'pdf':
                    path = default_storage.save(str(rand) + '.pdf', ContentFile(files.read()))
                    print(BASE_DIR)
                    pdf_file = (str(BASE_DIR)+'/media/') + str(path)
                    docx_file = (str(BASE_DIR)+'/media/') + (str(rand)+'.docx')
                    global texts

                    def convert_pdf_to_docx():
                        cv = Converter(pdf_file)
                        cv.convert(docx_file)
                        cv.close()

                    def read_doc_save_var():
                        convert_pdf_to_docx()
                        doc = Document(docx_file)
                        texts = ""
                        for para in doc.paragraphs:
                            texts += para.text
                        return texts

                    text = read_doc_save_var()
                    context = {'file': str(self.i)+'.txt', 'Abstract': self.process(text)}
                    return Response(self.best(text), status=status.HTTP_200_OK)

                elif str(filename_split[1]) == 'docx':
                    path = default_storage.save(str(rand) + '.docx', ContentFile(files.read()))
                    docx_file = (str(BASE_DIR) + '/media/') + str(path)
                    global texts
                    def read_doc_save():
                        doc = Document(docx_file)
                        texts = ""
                        for para in doc.paragraphs:
                            texts += para.text
                        return texts
                    text = read_doc_save()
                    context = {'file': str(self.i)+'.txt', 'Abstract': self.process(text)}
                    return Response(self.best(text), status=status.HTTP_200_OK)

                elif str(filename_split[1]) == 'txt':
                    path = default_storage.save(str(rand) + '.txt', ContentFile(files.read()))
                    doc_file = (str(BASE_DIR) + '/media/') + str(path)
                    with open(doc_file, 'r') as file:
                        text = file.read()
                    context = {'file': str(self.i)+'.txt', 'Abstract': self.process(text)}
                    return Response(self.best(text), status=status.HTTP_200_OK)
                return Response({'msg':'Yes it has a '+filename_split[1]+' file'},status=status.HTTP_200_OK)
            else:
                return Response({'msg':'please upload "docx","pdf","txt"'},status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response({'msg':'please upload file and submit'},status=status.HTTP_400_BAD_REQUEST)